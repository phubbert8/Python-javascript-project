# This is a Python script to
# create a CV document.

from docx import Document
from docx.shared import Inches
import pyttsx3


def speak(text):
    pyttsx3.speak(text)


document = Document()
document.add_picture('ME.JPG', width=Inches(1.5))

# pic = Image.open('ME.JPG');
# resizeImage = pic.resize(round(pic.size[0]*.5), round(pic.size[1]*.5)));
# rotate = resizeImage.rotate(270);
# resizeImage.show();

# name phone and email
name = input('What is your name? ')
speak('hello ' + name + ' How are you today?')
phone_number = input('What is your phone number? ')
email = input('What is your email? ')

document.add_paragraph(name + '|' + phone_number + '|' + email)

#  about me
document.add_heading('About Me')
document.add_paragraph(input('Tell about yourself? '))
# school experience
document.add_heading('Education')
p = document.add_paragraph()

school = input('Enter school ')
start_date = input('Start Date ')
grad_date = input('Enter date of grad ')

p.add_run(school + ' ').bold = True
p.add_run(start_date + '-' + grad_date + '\n').italic = True

# more school

while True:
    has_more_schools = input(
        'Do you have more schools? Yes or No ')
    if has_more_schools.lower() == 'yes':
        p = document.add_paragraph()

        school = input('Enter school ')
        start_date = input('Start Date ')
        grad_date = input('Enter date of grad ')

        p.add_run(school + ' ').bold = True
        p.add_run(start_date + '-' + grad_date + '\n').italic = True
    else:
        break

# work experience
document.add_heading('Work Experience')
p = document.add_paragraph()

company = input('Enter company ')
from_date = input('From Date ')
to_date = input('To Date ')

p.add_run(company + ' ').bold = True
p.add_run(from_date + '-' + to_date + '\n').italic = True

experience_details = input(
    'Describe your experience at ' + company)
p.add_run(experience_details)

# more experiences
while True:
    has_more_experiences = input(
        'Do you have more experiences? Yes or No ')
    if has_more_experiences.lower() == 'yes':
        p = document.add_paragraph()

        company = input('Enter company ')
        from_date = input('From Date ')
        to_date = input('To Date ')

        p.add_run(company + ' ').bold = True
        p.add_run(from_date + '-' + to_date + '\n').italic = True

        experience_details = input(
            'Describe your experience at ' + company)
        p.add_run(experience_details)
    else:
        break
# Skills sections
document.add_heading('Skills')
skill = input('Enter skill ')
p = document.add_paragraph(skill)
p.style = 'List Bullet'

# more experiences
while True:
    have_more_skills = input(
        'Do you have more skills? Yes or No ')
    if have_more_skills.lower() == 'yes':
        skill = input('Enter skill ')
        p = document.add_paragraph(skill)
        p.style = 'List Bullet'
    else:
        break

# footer
section = document.sections[0]
footer = section.footer
p = footer.paragraphs[0]
p.text = "CV generated using a course with an Amigoscode's project"

document.save('cv.docx')
